"""
Resume Depth - resume analyzer.

Frontend (static/index.html) uploads a resume -> this server extracts the text ->
an LLM scores 4 dimensions -> results stream back to the browser one card at a
time.

The model is pluggable via .env:
    LLM_BACKEND=ollama   local, private, slow  (llama3.1:8b)
    LLM_BACKEND=groq     hosted, fast, metered (see GROQ_MODEL)

Run:  uvicorn app:app --reload --port 8000
"""

import asyncio
import io
import json
import os
import re
from datetime import datetime
from pathlib import Path

import httpx
from dotenv import load_dotenv
from fastapi import FastAPI, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles

BASE_DIR = Path(__file__).parent
load_dotenv(BASE_DIR / ".env")

# "groq" = hosted and fast; "ollama" = local and private. See .env.
LLM_BACKEND = os.getenv("LLM_BACKEND", "ollama").strip().lower()

# RunPod's console shows the pod's OpenAI-compatible URL, which ends in /v1, and
# pasting that straight in takes the whole app down: every call here goes to the
# native API at /api/chat (the /v1 layer cannot set num_ctx), so the suffix turns
# each one into /v1/api/chat and a 404. Strip it rather than fail on it.
OLLAMA_URL = re.sub(r"/v1/?$", "",
                    os.getenv("OLLAMA_URL", "http://localhost:11434").strip()).rstrip("/")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3.1:8b")

GROQ_URL = os.getenv("GROQ_URL", "https://api.groq.com/openai/v1").rstrip("/")
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_MODEL = os.getenv("GROQ_MODEL", "openai/gpt-oss-120b")

MODEL = GROQ_MODEL if LLM_BACKEND == "groq" else OLLAMA_MODEL

# Ollama can now point at a remote GPU, so the backend name no longer tells us
# whether the resume stays on this machine. Only the host does.
IS_LOCAL = LLM_BACKEND != "groq" and bool(
    re.match(r"https?://(localhost|127\.0\.0\.1|\[::1\])(:|/|$)", OLLAMA_URL))
HOST = re.sub(r"^https?://", "", GROQ_URL if LLM_BACKEND == "groq" else OLLAMA_URL).split("/")[0]
NUM_CTX = 16384              # ollama only: must hold resume + prompt + reply

# Local Ollama only cares about context size, so it can read a whole 8-page CV.
# Groq's free tier meters tokens per minute (8000 on the gpt-oss models), and a
# 28k-char resume alone is ~7k tokens - one request would blow the whole minute.
MAX_RESUME_CHARS = int(os.getenv(
    "GROQ_MAX_RESUME_CHARS" if LLM_BACKEND == "groq" else "OLLAMA_MAX_RESUME_CHARS",
    "12000" if LLM_BACKEND == "groq" else "28000"))

# The frontend is no longer necessarily served by this process - it can sit on
# its own host and call this one across origins. Every origin allowed to do that
# is listed in ALLOWED_ORIGINS, comma separated. Empty means same-origin only,
# which is what a local `uvicorn app:app` run wants.
ALLOWED_ORIGINS = [o.strip().rstrip("/")
                   for o in os.getenv("ALLOWED_ORIGINS", "").split(",") if o.strip()]

app = FastAPI(title="Resume Depth")

if ALLOWED_ORIGINS:
    app.add_middleware(
        CORSMiddleware,
        allow_origins=ALLOWED_ORIGINS,
        allow_methods=["GET", "POST", "OPTIONS"],
        allow_headers=["*"],
    )

app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")


# --------------------------------------------------------------------------
# 1. Text extraction
# --------------------------------------------------------------------------

def extract_text(filename: str, blob: bytes) -> str:
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(blob))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    elif ext == ".docx":
        import docx
        doc = docx.Document(io.BytesIO(blob))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        text = "\n".join(parts)
    elif ext in (".txt", ".md"):
        text = blob.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type '{ext}'. Use PDF, DOCX, TXT or MD.")

    # squeeze blank lines / trailing spaces so the model sees clean structure
    text = "\n".join(line.rstrip() for line in text.splitlines())
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if not text:
        raise ValueError("No text found in this file (is it a scanned image PDF?).")

    # A plain head-truncation on a long CV throws away Education and the earliest
    # jobs, which is exactly what Completeness and Growth need. Keep both ends.
    if len(text) > MAX_RESUME_CHARS:
        head = int(MAX_RESUME_CHARS * 0.65)
        tail = MAX_RESUME_CHARS - head
        text = text[:head] + "\n\n[... middle of resume omitted ...]\n\n" + text[-tail:]
    return text


# --------------------------------------------------------------------------
# 1b. Is this actually a resume?
# --------------------------------------------------------------------------

# Every check below assumes a CV. Hand it a technical report and it answers
# anyway - scoring a LoRA merge note as "Junior Dev (no start date) -> Technical
# Specialist (2023)", with `safe_merge` and `torch_dtype` counted as claimed
# skills. Nothing in the report is wrong to the model; it was just never asked
# whether the document was a resume at all.
#
# Resumes carry marks that reports do not: a way to reach the person, and
# headed sections. Measured over 6 real resumes and 13 real reports, every
# resume scored 6 or more and no report scored above 2, so the line sits at 4 -
# clear of both, and biased towards letting an odd resume through rather than
# turning one away.

CONTACT_EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[a-z]{2,}", re.I)
CONTACT_PHONE = re.compile(r"(?:\+?\d{1,3}[\s.-]?)?\(?\d{3,5}\)?[\s.-]?\d{3}[\s.-]?\d{3,4}")
CONTACT_LINK = re.compile(r"linkedin\.com|github\.com", re.I)
RESUME_SECTION = re.compile(
    r"^\s*(work experience|professional experience|experience|employment|education|"
    r"academic|skills|technical skills|projects|certifications?|achievements|"
    r"summary|profile|objective|internships?|languages|interests|declaration)\s*:?\s*$",
    re.I | re.M)
RESUME_DATES = re.compile(
    r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}\b"
    r"|\b(19|20)\d{2}\s*[-\u2013\u2014]\s*((19|20)\d{2}|present|current)\b", re.I)

RESUME_MIN_SIGNAL = 4


def resume_signal(text: str) -> int:
    """How strongly this document reads as somebody's CV."""
    score = 0
    if CONTACT_EMAIL.search(text):
        score += 2                                   # the single clearest mark
    if CONTACT_PHONE.search(text):
        score += 1
    if CONTACT_LINK.search(text):
        score += 1
    score += min(3, len(RESUME_SECTION.findall(text)))
    if RESUME_DATES.search(text):
        score += 1
    return score


def require_resume(text: str) -> None:
    """Raise if this is not a resume, so analyse never runs on a report."""
    if resume_signal(text) >= RESUME_MIN_SIGNAL:
        return
    raise ValueError(
        "This does not look like a resume. It has no contact details and none of "
        "the sections a CV has - experience, education, skills. Upload a resume "
        "and the four checks will have something to read."
    )

# --------------------------------------------------------------------------
# 2. Ollama plumbing
# --------------------------------------------------------------------------

# Generation can genuinely take many minutes for an 8B model on CPU, and Ollama
# queues requests, so a short read timeout just turns a slow answer into a crash.
LLM_TIMEOUT = httpx.Timeout(connect=10.0, read=1800.0, write=60.0, pool=1800.0)


async def call_ollama(client: httpx.AsyncClient, system: str, user: str,
                      num_predict: int) -> str:
    resp = await client.post(
        f"{OLLAMA_URL}/api/chat",
        json={
            "model": OLLAMA_MODEL,
            "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            "format": "json",
            "stream": False,
            "options": {
                "temperature": 0.1,
                "num_ctx": NUM_CTX,
                "num_predict": num_predict,
            },
        },
        timeout=LLM_TIMEOUT,
    )
    resp.raise_for_status()
    return resp.json()["message"]["content"]


async def call_groq(client: httpx.AsyncClient, system: str, user: str,
                    num_predict: int) -> str:
    resp = await client.post(
        f"{GROQ_URL}/chat/completions",
        headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
        json={
            "model": GROQ_MODEL,
            "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            "response_format": {"type": "json_object"},
            "temperature": 0.1,
            # reasoning models spend part of this budget before the JSON starts
            "max_tokens": max(num_predict * 2, 4096),
        },
        timeout=LLM_TIMEOUT,
    )
    if resp.status_code == 429:                          # free tier is metered per minute
        raise RateLimited(float(resp.headers.get("retry-after", 20)))
    if resp.status_code >= 400:                          # surface Groq's own message
        try:
            detail = resp.json()["error"]["message"]
        except Exception:
            detail = resp.text[:200]
        raise RuntimeError(f"Groq {resp.status_code}: {detail}")
    return resp.json()["choices"][0]["message"]["content"]


class RateLimited(Exception):
    def __init__(self, wait: float):
        self.wait = min(wait, 65.0)
        super().__init__(f"rate limited, waiting {self.wait:.0f}s")


async def ask_llm(client: httpx.AsyncClient, system: str, user: str,
                  num_predict: int = 1600) -> dict:
    """One JSON-mode call to whichever backend is configured. Returns a dict."""
    call = call_groq if LLM_BACKEND == "groq" else call_ollama
    last_error = None
    for attempt in range(4):
        try:
            return parse_json(await call(client, system, user, num_predict))
        except RateLimited as limit:
            last_error = limit                           # the minute window resets
            await asyncio.sleep(limit.wait)
        except Exception as exc:
            last_error = exc                             # one retry; models drift
            if attempt >= 1:
                break
    raise describe(last_error)


def parse_json(content: str) -> dict:
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", content, re.S)      # model wrapped it in prose
        if match:
            return json.loads(match.group(0))
        raise ValueError(f"Model did not return JSON: {content[:200]}")


def describe(exc: Exception) -> Exception:
    """httpx timeouts stringify to '', which reaches the UI as a blank error."""
    if str(exc):
        return exc
    return RuntimeError(f"{type(exc).__name__} (the model took too long to answer)")


def clamp_score(value) -> int:
    try:
        n = int(round(float(value)))
    except (TypeError, ValueError):
        return 50
    return max(0, min(100, n))


def status_for(score: int) -> str:
    if score >= 70:
        return "pass"
    if score >= 50:
        return "warn"
    return "fail"


# 8B models happily invent skills the resume never claimed, which poisons both the
# chips and the ratio they are scored on. Keep only skills whose name literally
# appears in the resume, then compute the score ourselves from what survived, so
# the number always matches the chips on screen.
SKILL_WEIGHT = {"strong": 1.0, "moderate": 0.6, "claimed": 0.15}


def ground_evidence(data: dict, resume: str) -> dict:
    skills = data.get("skills")
    if not isinstance(skills, list):
        return data

    haystack = resume.lower()
    kept, seen = [], set()
    for item in skills:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name") or "").strip()
        key = name.lower()
        # CVs that write their skills as prose bullets tempt the model into
        # returning a whole sentence as one "skill" - a real one is a short name.
        if not name or len(name) > 40 or key in seen or key not in haystack:
            continue
        seen.add(key)
        strength = item.get("strength")
        kept.append({
            "name": name,
            "strength": strength if strength in SKILL_WEIGHT else "moderate",
            "note": str(item.get("note") or ""),
        })

    if not kept:
        return data

    proven = sum(1 for s in kept if s["strength"] != "claimed")
    data["skills"] = kept
    data["score"] = round(100 * sum(SKILL_WEIGHT[s["strength"]] for s in kept) / len(kept))
    data["detail"] = f"{proven} of {len(kept)} claimed skills appear in the described work"
    return data


DEGREE_SYNONYMS = {"bachelor", "bachelors", "b.tech", "btech", "b.e", "be", "b.s", "bs", "b.sc", "bsc", "degree"}

def ground_completeness(data: dict, resume: str) -> dict:
    """The model also calls things missing that are plainly in the document.

    A 'missing' item names the thing it is looking for. Take its proper nouns
    and check them against the text. Also handle degree equivalences (B.Tech = Bachelor's)
    and discard false positive claims when Education is already present.
    """
    missing = data.get("missing")
    if not isinstance(missing, list):
        return data

    text = resume.lower()
    # An explicit "present": null reaches here as None, which .get's default
    # never covers - and an exception in a check costs the whole card.
    present = [str(p).lower() for p in (data.get("present") or [])]
    has_education_present = any("education" in p or "degree" in p for p in present)
    has_degree_in_text = any(d in text for d in ("b.tech", "btech", "b.e", "b.s", "b.sc", "bachelor", "degree"))

    kept = []
    for item in missing:
        label = str(item).strip()
        if not label:
            continue
        label_lower = label.lower()

        # If education/degree is present in resume text or 'present' list:
        if has_education_present or has_degree_in_text:
            if any(d in label_lower for d in ("education", "degree", "bachelor", "master")):
                # Extract key field/topic words from missing label (e.g. artificial, intelligence, data, science)
                words_in_label = [w for w in re.findall(r"\b[a-z]{3,}\b", label_lower)
                                  if w not in ("education", "degree", "bachelor", "bachelors", "relevant", "field", "such", "position", "expected", "required", "with", "lack", "don't", "have")]
                matched_words = sum(1 for w in words_in_label if w in text)
                if not words_in_label or (len(words_in_label) > 0 and matched_words / len(words_in_label) >= 0.4):
                    continue  # Discard false positive education/degree claim!

        names = re.findall(r"\b[A-Z][A-Za-z0-9+#.\-]{2,}\b", label)
        names_to_check = [n for n in names if n.lower() not in DEGREE_SYNONYMS and n.lower() not in ("education", "relevant", "bachelor", "bachelors", "degree")]
        if names_to_check and all(n.lower() in text for n in names_to_check):
            continue  # provably present

        kept.append(label)

    data["missing"] = kept

    # The model picked its score believing those items were absent, so removing
    # them without touching the number leaves the card contradicting itself: a
    # 70 sitting above a list that reads 80. Recompute from what survived - but
    # only when something was actually removed. With nothing removed the model's
    # number stands, because it weighs a missing summary against a missing
    # portfolio in a way this ratio cannot, and because a model that returns a
    # short "missing" list would otherwise be handed a free 100.
    removed = len(missing) - len(kept)
    total = len(kept) + len(data.get("present") or [])
    if removed and total:
        found = total - len(kept)
        data["score"] = round(100 * found / total)
        data["detail"] = f"{found} of {total} expected sections present"

    # If missing items were cleaned up, adjust reasoning if it contains false degree claims
    reasoning = data.get("reasoning", "")
    if reasoning and ("bachelor" in reasoning.lower() or "degree" in reasoning.lower()) and has_degree_in_text:
        # Remove sentences that falsely claim lack of degree when B.Tech/B.S is present
        sentences = re.split(r"(?<=[.!?])\s+", reasoning)
        cleaned_sentences = [
            s for s in sentences
            if not (any(d in s.lower() for d in ("bachelor", "degree", "education")) and any(neg in s.lower() for neg in ("don't have", "lack", "missing", "without", "no bachelor")))
        ]
        data["reasoning"] = " ".join(cleaned_sentences) if cleaned_sentences else reasoning

    return data


# Resumes are written newest-first, and an 8B model reads one top-to-bottom and
# calls that "date order". The progression then comes back narrated backwards -
# "you started as an Intern in Jan 2026, then became Project Head in Apr 2025".
# Same fix as the skills above: let the model name the roles, but do the sorting
# in Python, where a date is a date.

MONTHS = {"jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
          "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12}
# Internships are often dated by season rather than by month.
SEASONS = {"winter": 1, "spring": 3, "summer": 6, "fall": 9, "autumn": 9}
PRESENT_WORDS = ("present", "current", "now", "ongoing", "to date", "today")


def parse_when(text) -> tuple | None:
    """A date as a resume writes it -> a sortable (year, month).

    Handles 'Jan 2026', 'January 2026', '01/2026', '2026-01', 'Summer 2024',
    a bare '2026', and ranges like 'Jan 2020 - Mar 2022' (the start date wins).
    Returns None when there is no year to anchor on, because an undated role
    cannot be placed on the timeline at all.
    """
    if not text:
        return None
    s = str(text).strip().lower()
    if not re.search(r"\d{4}", s):
        # "Present" with no year of its own belongs at the end.
        return (9999, 12) if any(w in s for w in PRESENT_WORDS) else None

    # Numeric dates first: their own separator is a hyphen, so splitting the
    # range before reading them would leave "2026-01" looking like a bare year.
    m = re.match(r"^(\d{4})[-/](\d{1,2})\b", s)          # 2026-01
    if m:
        return (int(m.group(1)), max(1, min(12, int(m.group(2)))))
    m = re.match(r"^(\d{1,2})[-/](\d{4})\b", s)          # 01/2026
    if m:
        return (int(m.group(2)), max(1, min(12, int(m.group(1)))))

    # A range carries two dates; a role starts at the first of them.
    s = re.split(r"\s*(?:-|–|—|to|until|through)\s*", s)[0].strip()

    year = re.search(r"\b(?:19|20)\d{2}\b", s)
    if not year:
        return None

    month = 0                                            # year only sorts first
    for word in re.findall(r"[a-z]+", s):
        if word in SEASONS:
            month = SEASONS[word]
            break
        if word[:3] in MONTHS:
            month = MONTHS[word[:3]]
            break
    return (int(year.group(0)), month)


# "Title (Jan 2026)" pairs, for models that ignore the roles array and only
# fill in the detail line. Arrows and commas end a title, never start one.
ROLE_IN_DETAIL = re.compile(r"([^()→>,;]+?)\s*\(([^()]*\d{4}[^()]*)\)")


def in_date_order(roles: list) -> bool:
    """Dates, not titles - two spells under one title are still orderable."""
    return all(roles[i][2] <= roles[i + 1][2] for i in range(len(roles) - 1))


def roles_from_detail(data: dict) -> list:
    """The "Title (Jan 2026)" pairs written into the detail line."""
    return [(title.strip(" -→\t"), when.strip(), parse_when(when))
            for title, when in ROLE_IN_DETAIL.findall(str(data.get("detail") or ""))]


def collect_roles(data: dict) -> list:
    """Every role the model named, as (title, date-as-written, sort key)."""
    roles = []
    for item in data.get("roles") or []:
        if not isinstance(item, dict):
            continue
        title = str(item.get("title") or "").strip()
        when = str(item.get("start") or item.get("date") or "").strip()
        if title:
            roles.append((title, when, parse_when(when)))

    if len(roles) < 2:      # the model skipped the array - read the detail line
        roles = roles_from_detail(data)
    return roles


def ground_growth(data: dict, resume: str) -> dict:
    """Put the roles on the detail line back in date order.

    This touches the one-line summary and nothing else. The verdict, the
    reasoning and the trajectory are the model's read of the career and stay
    exactly as written - it saw what the person actually did, which is the part
    worth having. Sorting a list is the only thing here that Python does better.

    Undated roles are dropped: there is nowhere to put them on a timeline.
    """
    dated = [r for r in collect_roles(data) if r[2]]
    if len(dated) < 2:
        return data

    shown = [r for r in roles_from_detail(data) if r[2]]
    if in_date_order(dated) and (len(shown) < 2 or in_date_order(shown)):
        return data                       # already chronological, leave it alone

    ordered = sorted(dated, key=lambda role: role[2])
    data["detail"] = " → ".join(f"{title} ({when})" for title, when, _ in ordered)
    return data


# Tenure is counted here, not by the model. Asked for "total_years" it returns an
# impression - and it counts side projects and coursework as experience, which is
# exactly what a recruiter does not want. Employment is a set of dated intervals;
# adding them up is arithmetic.

INTERNSHIP_WORDS = re.compile(
    r"\b(intern|internship|trainee|apprentice|co-?op|industrial training|"
    r"summer (?:analyst|associate)|student)\b", re.I)

MONTHS_PER_YEAR = 12.0


def month_index(when: tuple) -> int:
    """(year, month) -> a single number, so intervals can be compared."""
    year, month = when
    return year * 12 + max(1, month) - 1


def parse_end(text, started: tuple | None) -> tuple | None:
    """The end of a spell. 'Present' means now; a range means its second date."""
    if not text:
        return None
    body = str(text).strip().lower()
    if any(word in body for word in PRESENT_WORDS):
        today = datetime.now()
        return (today.year, today.month)
    # "Jan 2020 - Mar 2022" hands us the whole range; the end is what follows.
    parts = re.split(r"\s*(?:-|\u2013|\u2014|to|until|through)\s*", body)
    when = parse_when(parts[-1] if len(parts) > 1 else body)
    if when == (9999, 12):                       # an open-ended spell runs to now
        today = datetime.now()
        return (today.year, today.month)
    return when


def merge_months(spans: list) -> int:
    """Total months covered by these spans, counting overlap only once.

    Two jobs held at the same time are one stretch of a career, not two - a
    naive sum would hand someone eight years for four years of moonlighting.
    """
    if not spans:
        return 0
    total, current_start, current_end = 0, *spans[0]
    for start, end in sorted(spans)[1:]:
        if start <= current_end:                 # overlapping or touching
            current_end = max(current_end, end)
        else:
            total += current_end - current_start
            current_start, current_end = start, end
    return total + current_end - current_start


def say_duration(months: int) -> str:
    if months <= 0:
        return "none"
    if months < 12:
        return f"{months} month{'s' if months != 1 else ''}"
    years = months / MONTHS_PER_YEAR
    return f"{years:.1f} years".replace(".0 years", " years")


def collect_positions(data: dict) -> list:
    """The jobs the model listed, dated and classified. Projects never reach here."""
    positions = []
    for item in data.get("roles") or []:
        if not isinstance(item, dict):
            continue
        title = str(item.get("title") or "").strip()
        if not title:
            continue
        employer = str(item.get("employer") or "").strip()
        start_raw = str(item.get("start") or "").strip()
        end_raw = str(item.get("end") or "").strip()

        started = parse_when(start_raw)
        ended = parse_end(end_raw, started) if end_raw else None
        if started and started == (9999, 12):    # a start of "Present" is nonsense
            started = None

        # Trust the model's label, but a title that plainly says intern wins.
        kind = str(item.get("type") or "").strip().lower()
        if INTERNSHIP_WORDS.search(f"{title} {employer}"):
            kind = "internship"
        elif kind not in ("work", "internship"):
            kind = "work"

        months = 0
        if started and ended and month_index(ended) > month_index(started):
            months = month_index(ended) - month_index(started)

        positions.append({
            "title": title,
            "employer": employer,
            "start": start_raw,
            "end": end_raw or ("Present" if started and not end_raw else ""),
            "type": kind,
            "months": months,
            "_span": (month_index(started), month_index(ended))
            if started and ended and month_index(ended) > month_index(started) else None,
        })

    positions.sort(key=lambda p: p["_span"][0] if p["_span"] else 10 ** 9)
    return positions


def ground_experience(data: dict, resume: str) -> dict:
    """Count tenure from the dated positions, split work from internships.

    Projects are excluded because the prompt asks only for positions, and the
    two totals are kept apart because six months of interning is not six months
    of being employed to do the job.
    """
    positions = collect_positions(data)
    if not positions:
        return data

    work = [p["_span"] for p in positions if p["type"] == "work" and p["_span"]]
    intern = [p["_span"] for p in positions if p["type"] == "internship" and p["_span"]]

    work_months, intern_months = merge_months(work), merge_months(intern)
    data["total_years"] = say_duration(work_months) if work_months else "no dated roles"
    data["internship_years"] = say_duration(intern_months)
    data["roles"] = [{k: v for k, v in p.items() if k != "_span"} for p in positions]
    return data

# --------------------------------------------------------------------------
# 3. The four checks
# --------------------------------------------------------------------------

SYSTEM = (
    "You are a strict technical recruiter reviewing a resume. "
    "You judge only what is written in the resume - never invent facts. "
    "Never write 'I' - address the candidate as 'you'. "
    "You reply with JSON only, no prose, no markdown fences."
)

ROLE_PROMPT = """Read the resume and identify what job the candidate is positioned for.

RESUME:
{resume}

Reply with JSON exactly like:
{{"role": "<job title, e.g. Backend Engineer>",
  "seniority": "<Intern|Junior|Mid-level|Senior|Lead|Manager>",
  "expected_sections": ["<section a resume for THIS role must have>", "..."]}}

Base "expected_sections" on what this specific role and seniority is judged on
(e.g. a data scientist needs projects with datasets and metrics; a manager needs
team size and scope; a fresher needs education and academic projects)."""


CHECKS = [
    {
        "key": "experience",
        "title": "Experience",
        "why": "The system reads the project content to see what work they have really "
               "done, and the dates to see how long they did it for.",
        "prompt": """Judge the DEPTH OF EXPERIENCE in this resume.

Detected role: {role} ({seniority})

RESUME:
{resume}

Rules:
- Read the actual project/job descriptions. Real engineering work (systems built,
  problems solved, decisions made) scores high. Vague duty lists score low.
- Read the dates. Long tenure and continuous history score high. Gaps, very short
  stints, or missing dates score low.
- No dates anywhere = cap the score at 55.

Reply with JSON:
{{"score": <0-100>,
  "verdict": "<one sentence addressed to the candidate as 'Your experience ...'>",
  "detail": "<one short sentence naming the concrete evidence you used>",
  "reasoning": "<3-4 sentences. Name the actual employers, projects and dates you
    read. Say what the work shows about their depth, and what specifically held the
    score back. Address the candidate as 'you'. Quote real phrases from the resume.>",
  "roles": [{{"title": "<job title>",
             "employer": "<company or organisation>",
             "start": "<start date exactly as the resume writes it>",
             "end": "<end date as written, or Present>",
             "type": "<work|internship>"}}]}}

"roles" is every POSITION held, oldest first - a job someone employed them to do.
Internships, trainee and co-op placements are positions too; mark those
"internship". Personal projects, college coursework, hackathons and portfolio
sites are NOT positions and must not appear, however impressive. Copy the dates
exactly as written; leave a date empty rather than inventing one.""",
    },
    {
        "key": "evidence",
        "title": "Evidence",
        "why": "We check that the skills they claim actually show up in the work they described.",
        "prompt": """Check whether the CLAIMED SKILLS are backed by the described work.

Detected role: {role} ({seniority})

RESUME:
{resume}

Steps:
1. List every skill/tool/technology the resume claims (skills section, summary, anywhere).
   Copy the names EXACTLY as written in the resume. Never add a technology that
   does not appear in the text above - a wrong name invalidates the whole check.
2. For each one, look for it inside the experience or project descriptions.
   - "strong"   = used in a described project WITH context or an outcome
   - "moderate" = mentioned in a project but with no detail
   - "claimed"  = listed only in the skills section, never used anywhere

Score = how much of the skill list is actually proven. Mostly "claimed" scores low.

Reply with JSON:
{{"score": <0-100>,
  "verdict": "<one sentence addressed to the candidate>",
  "detail": "<one short sentence, e.g. '9 of 14 claimed skills appear in project work'>",
  "reasoning": "<3-4 sentences. Name which skills are genuinely proven and where
    they were proven. Then name the ones that appear only in the skills list, and
    say what evidence would fix them. Address the candidate as 'you'.>",
  "skills": [{{"name": "<skill>", "strength": "strong|moderate|claimed",
               "note": "<where it was proven, or 'not used in any project'>"}}]}}

Include every skill you found, up to 20.""",
    },
    {
        "key": "growth",
        "title": "Growth",
        "why": "Do the job titles rise over time? Junior to Senior to Lead shows someone "
               "trusted with more. A flat line for eight years does not.",
        "prompt": """How has this candidate GROWN over the course of their career?

Detected role: {role} ({seniority})

RESUME:
{resume}

Resumes are written newest-first, so read the dates rather than the page order.
Judge what actually changed from the earliest position to the latest: the work
they were trusted with, the scope they owned, the difficulty of what they built.
Whatever shape this particular career has, judge growth on its own terms - a run
of internships, a switch of field, one long job and a person who has only ever
been promoted are four different stories, and each has its own kind of progress.

Reply with JSON:
{{"score": <0-100>,
  "roles": [{{"title": "<the position held - a job, not a project>",
             "start": "<its start date, exactly as the resume writes it>"}}],
  "verdict": "<one sentence to the candidate on how their career has developed>",
  "detail": "<the progression in a few words, e.g. 'Junior Dev (2019) -> Senior Dev (2023)'>",
  "reasoning": "<3-4 sentences on how they grew and what tells you so - what they
    were doing early on, what they are trusted with now, and what that says about
    them. Point at the evidence rather than restating the timeline. Address the
    candidate as 'you'.>",
  "trajectory": "<rising|steady|flat|unclear>"}}""",
    },
    {
        "key": "completeness",
        "title": "Completeness",
        "why": "Are the basic parts there - a summary, a skill list, education, certificates. "
               "Missing pieces cost the candidate nothing to fix.",
        "prompt": """Check whether this resume has all the parts expected FOR ITS ROLE.

Detected role: {role} ({seniority})
Sections this role is expected to have: {expected_sections}

RESUME:
{resume}

Rules for evaluation:
- Universals to check: contact details, professional summary, skills list, work experience with dates, education.
- Degree Equivalence: B.Tech, B.E., B.S., B.Sc., M.Tech, M.S. are valid degrees. If the candidate has "B.Tech in Artificial Intelligence and Data Science", they DO have a Bachelor's degree in AI & Data Science. Never say their degree is missing!
- Do NOT list Education or a degree as missing if an Education section with a relevant degree exists in the resume.

Score = share of expected parts that are present and usable.

Reply with JSON:
{{"score": <0-100>,
  "verdict": "<one sentence addressed to the candidate>",
  "detail": "<one short sentence>",
  "reasoning": "<3-4 sentences. Say which expected sections you found and where.
    Then name what is absent and why it matters for THIS role specifically.
    Address the candidate as 'you'.>",
  "present": ["<section found>", "..."],
  "missing": ["<expected section that is absent>", "..."]}}""",
    },
]


ACTION_PROMPT = """You reviewed a resume for a {role} ({seniority}) and scored it:

{summary}

Write the single highest-value fix for this candidate - start with the weakest
area. Two or three sentences, plain language, addressed to the candidate as "you".
Say what to do, not what is wrong.

Reply with JSON: {{"action": "<the advice>"}}"""


# --------------------------------------------------------------------------
# 4. Streaming endpoint
# --------------------------------------------------------------------------

def sse(event: str, data: dict) -> str:
    return f"event: {event}\ndata: {json.dumps(data)}\n\n"


async def preflight(client: httpx.AsyncClient) -> str:
    """Fail with a useful sentence before promising the user a report."""
    if LLM_BACKEND == "groq":
        if not GROQ_API_KEY:
            return "GROQ_API_KEY is not set. Add it to the .env file next to app.py."
        try:
            resp = await client.get(f"{GROQ_URL}/models",
                                    headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
                                    timeout=15.0)
        except Exception as exc:
            return f"Could not reach Groq: {describe(exc)}"
        if resp.status_code == 401:
            return "Groq rejected the API key in .env (401). Check GROQ_API_KEY."
        if resp.status_code >= 400:
            return f"Groq returned {resp.status_code}: {resp.text[:160]}"
        names = [m["id"] for m in resp.json().get("data", [])]
        if GROQ_MODEL not in names:
            return (f"Groq has no model '{GROQ_MODEL}'. Available: "
                    + ", ".join(sorted(names)[:8]))
        return ""

    try:
        tags = (await client.get(f"{OLLAMA_URL}/api/tags", timeout=10.0)).json()
    except Exception:
        return f"Ollama is not reachable at {OLLAMA_URL}. Start it and try again."
    names = [m["name"] for m in tags.get("models", [])]
    if not any(n == OLLAMA_MODEL or n.startswith(OLLAMA_MODEL.split(":")[0]) for n in names):
        return f"Model '{OLLAMA_MODEL}' is not installed. Run:  ollama pull {OLLAMA_MODEL}"
    return ""


# A single check can run for minutes, and nothing is sent while it does. Once the
# browser is on another host the stream crosses proxies that cut a connection
# that has gone quiet, so fill the gaps. A comment line is valid SSE that no
# client turns into an event.
HEARTBEAT_SECONDS = 15
KEEPALIVE = ": keepalive\n\n"


async def pending(coro):
    """Run a coroutine, yielding a keepalive for every quiet interval.

    Yields KEEPALIVE strings while it waits and the finished Task last, so the
    caller both keeps the stream warm and gets the result (or the exception)
    without this having to know what either looks like.
    """
    task = asyncio.create_task(coro)
    while True:
        done, _ = await asyncio.wait({task}, timeout=HEARTBEAT_SECONDS)
        if done:
            break
        yield KEEPALIVE
    yield task


async def run_analysis(resume: str):
    async with httpx.AsyncClient() as client:
        problem = await preflight(client)
        if problem:
            yield sse("error", {"message": problem})
            return

        yield sse("step", {"message": "Reading the resume..."})
        async for beat in pending(
                ask_llm(client, SYSTEM, ROLE_PROMPT.format(resume=resume))):
            if isinstance(beat, str):
                yield beat
        try:
            role_info = beat.result()
        except Exception as exc:
            yield sse("error", {"message": f"Model call failed: {exc}"})
            return

        role = role_info.get("role") or "Professional"
        seniority = role_info.get("seniority") or "Mid-level"
        expected = role_info.get("expected_sections") or []
        yield sse("role", {"role": role, "seniority": seniority})

        ctx = {
            "resume": resume,
            "role": role,
            "seniority": seniority,
            "expected_sections": ", ".join(expected) if expected else "(use your own judgement)",
        }

        yield sse("step", {"message": f"Scoring against a {role} profile..."})

        # The positions Experience settles on, shared with Growth below.
        timeline = []

        async def run_check(check):
            # the skills array on a long CV is the one reply that can run long
            budget = 3000 if check["key"] == "evidence" else 1800
            data = await ask_llm(client, SYSTEM, check["prompt"].format(**ctx), budget)
            if check["key"] == "experience":
                data = ground_experience(data, resume)
                timeline[:] = data.get("roles") or []
            elif check["key"] == "evidence":
                data = ground_evidence(data, resume)
            elif check["key"] == "completeness":
                data = ground_completeness(data, resume)
            elif check["key"] == "growth":
                # Experience already worked the positions out and threw the
                # projects away. Growth showing a different list of the same
                # career is just two chances to be wrong, so it reuses that one.
                if timeline:
                    data["roles"] = [dict(role) for role in timeline]
                data = ground_growth(data, resume)
            score = clamp_score(data.get("score"))
            return {
                "key": check["key"],
                "title": check["title"],
                "why": check["why"],
                "score": score,
                "status": status_for(score),
                "verdict": data.get("verdict") or "",
                "detail": data.get("detail") or "",
                "reasoning": data.get("reasoning") or "",
                "extra": {k: v for k, v in data.items()
                          if k in ("skills", "present", "missing", "total_years",
                                       "internship_years", "roles", "trajectory")},
            }

        # Sequential on purpose. Ollama serialises requests anyway, so firing all
        # four at once buys no speed - it just leaves the last one sitting in the
        # queue until its read timeout fires.
        scored = []
        for check in CHECKS:
            async for beat in pending(run_check(check)):
                if isinstance(beat, str):
                    yield beat
            try:
                result = beat.result()
                scored.append(result)
            except Exception as exc:
                result = {
                    "key": check["key"],
                    "title": check["title"],
                    "why": check["why"],
                    "score": 0,
                    "status": "error",
                    "verdict": "This check could not be completed.",
                    "detail": str(describe(exc))[:180],
                    "reasoning": "",
                    "extra": {},
                }
            yield sse("check", result)

        if not scored:
            yield sse("error", {"message": "Every check failed - is Ollama still running?"})
            return

        overall = round(sum(r["score"] for r in scored) / len(scored))

        summary = "\n".join(f"- {r['title']}: {r['score']}/100 - {r['verdict']}" for r in scored)
        async for beat in pending(ask_llm(client, SYSTEM, ACTION_PROMPT.format(
                role=role, seniority=seniority, summary=summary), 400)):
            if isinstance(beat, str):
                yield beat
        try:
            action = beat.result().get("action", "")
        except Exception:
            action = ""

        yield sse("done", {"score": overall, "status": status_for(overall), "action": action})


@app.post("/api/analyze")
async def analyze(file: UploadFile = File(...)):
    blob = await file.read()
    try:
        resume = extract_text(file.filename or "resume.pdf", blob)
        require_resume(resume)
    except Exception as exc:
        message = str(exc)

        async def fail():
            yield sse("error", {"message": message})

        return StreamingResponse(fail(), media_type="text/event-stream")

    return StreamingResponse(
        run_analysis(resume),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.get("/api/backend")
async def backend():
    return {"backend": LLM_BACKEND, "model": MODEL, "local": IS_LOCAL, "host": HOST}


@app.get("/")
async def index():
    # index.html links its assets relatively so the same file can be dropped on a
    # static host. Serving it from / would resolve those against /, so hand the
    # browser the /static/ copy instead of the file.
    return RedirectResponse("/static/index.html")
